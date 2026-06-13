from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT_DIR = Path(__file__).resolve().parent
OUT_DIR = PROJECT_DIR / "report_assets"
OUT_DIR.mkdir(exist_ok=True)
REPORT_PATH = PROJECT_DIR / "NLP_Mini_Project_1_Final_Report_Human_Style.docx"

METRICS = {
    "articles": 5,
    "corpus_chars": 96199,
    "tokenizer": "khmernltk.word_tokenize",
    "total_tokens": 19681,
    "train_tokens": 13776,
    "valid_tokens": 1968,
    "test_tokens": 3937,
    "vocab_size": 500,
    "valid_unk": 660,
    "test_unk": 1433,
    "best_lambdas": "(0.4, 0.3, 0.2, 0.1)",
    "best_k": 0.001,
    "best_valid_pp": 47.4493,
    "lm1_test_pp": 27.2196,
    "lm2_test_pp": 45.7879,
}

VOCAB_ROWS = [
    (100, 100, 1164, 2344, 7.0151, 9.2468),
    (200, 200, 930, 2021, 11.3045, 16.1981),
    (300, 300, 808, 1783, 16.2645, 24.7400),
    (500, 500, 660, 1433, 27.2196, 45.7879),
    (700, 700, 574, 1233, 38.2488, 68.3284),
    (1000, 1000, 523, 1119, 49.2772, 92.2916),
]

SOURCE_URLS = [
    "https://km.wikipedia.org/wiki/ប្រាសាទ_អង្គរវត្ត",
    "https://km.wikipedia.org/wiki/អង្គរធំ",
    "https://km.wikipedia.org/wiki/ប្រាសាទបាយ័ន",
    "https://km.wikipedia.org/wiki/ខេត្តសៀមរាប",
    "https://km.wikipedia.org/wiki/ចក្រភពខ្មែរ",
]


def chart_paths():
    split_path = OUT_DIR / "human_dataset_split.png"
    model_path = OUT_DIR / "human_model_perplexity.png"
    vocab_path = OUT_DIR / "human_vocab_perplexity.png"

    plt.figure(figsize=(7, 4))
    labels = ["Training", "Validation", "Testing"]
    counts = [METRICS["train_tokens"], METRICS["valid_tokens"], METRICS["test_tokens"]]
    bars = plt.bar(labels, counts, color=["#2E86AB", "#F6C85F", "#6F4E7C"])
    plt.title("Corpus Split")
    plt.ylabel("Tokens")
    plt.grid(axis="y", alpha=0.25)
    for bar, count in zip(bars, counts):
        plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 120, f"{count:,}", ha="center")
    plt.tight_layout()
    plt.savefig(split_path, dpi=180)
    plt.close()

    plt.figure(figsize=(7, 4))
    bars = plt.bar(["LM1 Backoff", "LM2 Interpolation"], [METRICS["lm1_test_pp"], METRICS["lm2_test_pp"]], color=["#2E86AB", "#D1495B"])
    plt.title("Test Perplexity by Model")
    plt.ylabel("Perplexity (lower is better)")
    plt.grid(axis="y", alpha=0.25)
    for bar, value in zip(bars, [METRICS["lm1_test_pp"], METRICS["lm2_test_pp"]]):
        plt.text(bar.get_x() + bar.get_width() / 2, value + 0.8, f"{value:.4f}", ha="center")
    plt.tight_layout()
    plt.savefig(model_path, dpi=180)
    plt.close()

    caps = [row[0] for row in VOCAB_ROWS]
    lm1 = [row[4] for row in VOCAB_ROWS]
    lm2 = [row[5] for row in VOCAB_ROWS]
    plt.figure(figsize=(8, 4.5))
    plt.plot(caps, lm1, marker="o", linewidth=2, label="LM1 Backoff")
    plt.plot(caps, lm2, marker="o", linewidth=2, label="LM2 Interpolation")
    plt.title("Vocabulary Size vs Test Perplexity")
    plt.xlabel("Vocabulary cap")
    plt.ylabel("Test perplexity")
    plt.grid(alpha=0.25)
    plt.legend()
    plt.tight_layout()
    plt.savefig(vocab_path, dpi=180)
    plt.close()

    return split_path, model_path, vocab_path


def font_run(run, size=10.5, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans Khmer")
    run.font.size = Pt(size)
    run.bold = bold


def paragraph(doc, text="", size=10.5, bold=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    font_run(r, size=size, bold=bold)
    if align:
        p.alignment = align
    return p


def heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    font_run(r, size=14 if level == 1 else 12, bold=True)
    return p


def shade(cell, color="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        t.rows[0].cells[i].text = str(header)
        shade(t.rows[0].cells[i])
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                font_run(r, size=9.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    font_run(r, size=9)
    return t


def build():
    split_img, model_img, vocab_img = chart_paths()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mini Project 1: Text Generation with 4-gram Language Models")
    font_run(r, size=16, bold=True)
    paragraph(doc, "Natural Language Processing | Master of Data Science", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "Corpus topic: Khmer Wikipedia articles related to Angkor Wat", align=WD_ALIGN_PARAGRAPH.CENTER)

    heading(doc, "1. Introduction", 1)
    paragraph(
        doc,
        "The purpose of this mini project is to build a statistical language model that can learn word-sequence patterns from Khmer text and generate new text from a starting phrase. I selected Angkor Wat as the main topic because it has enough historical, cultural, and architectural vocabulary to make the generation task meaningful. Instead of using only one short article, I combined five related Khmer Wikipedia articles so that the model could learn from a larger and more consistent corpus.",
    )
    paragraph(
        doc,
        "The project compares two required 4-gram language models. The first model is an unsmoothed backoff model, and the second model is an interpolated model with add-k smoothing. Both models are evaluated using perplexity on a separate test set.",
    )

    heading(doc, "2. Corpus Collection and Preprocessing", 1)
    paragraph(
        doc,
        "The corpus was collected from five related Khmer Wikipedia pages: Angkor Wat, Angkor Thom, Bayon Temple, Siem Reap Province, and the Khmer Empire. These pages are related to the same historical and cultural domain, so combining them increases corpus size without mixing unrelated topics.",
    )
    table(
        doc,
        ["Corpus statistic", "Value"],
        [
            ["Number of articles", METRICS["articles"]],
            ["Corpus characters", f"{METRICS['corpus_chars']:,}"],
            ["Tokenizer", METRICS["tokenizer"]],
            ["Total tokens", f"{METRICS['total_tokens']:,}"],
            ["Training tokens", f"{METRICS['train_tokens']:,}"],
            ["Validation tokens", f"{METRICS['valid_tokens']:,}"],
            ["Testing tokens", f"{METRICS['test_tokens']:,}"],
            ["Vocabulary size used", METRICS["vocab_size"]],
        ],
    )
    paragraph(
        doc,
        "During preprocessing, I removed English-only headings and artificial article markers, separated Khmer sentence punctuation, and used khmernltk.word_tokenize to improve Khmer token quality. Tokens outside the limited vocabulary were replaced with <UNK>, which directly follows the project instruction.",
    )
    doc.add_picture(str(split_img), width=Inches(5.8))
    paragraph(
        doc,
        "The chart above confirms that the dataset follows the required 70/10/20 structure. The training set is used to learn n-gram counts, the validation set is used to tune LM2 hyperparameters, and the test set is kept for final evaluation.",
    )

    heading(doc, "3. Language Model Methods", 1)
    paragraph(
        doc,
        "A 4-gram model predicts the next token using the previous three tokens as context. For example, if the context is 'ប្រាសាទ អង្គរវត្ត ជា', the model estimates which token is most likely to come next based on patterns counted in the training data.",
    )
    paragraph(
        doc,
        "LM1 uses a backoff strategy. It first tries to use the full 4-gram probability. If that exact 4-gram was not observed in training, it backs off to trigram, then bigram, and finally unigram probability. This approach works well when exact local patterns exist in the corpus.",
    )
    paragraph(
        doc,
        "LM2 uses interpolation. Instead of choosing only one n-gram level, it combines unigram, bigram, trigram, and 4-gram probabilities using lambda weights. Add-k smoothing is also applied so unseen n-grams receive a small non-zero probability.",
    )

    heading(doc, "4. Hyperparameter Tuning", 1)
    paragraph(
        doc,
        "For LM2, I used grid search to tune lambda values and the smoothing value k. Each lambda combination controls how much weight is given to unigram, bigram, trigram, and 4-gram probabilities. The value k controls how much smoothing is applied. The best combination is selected by the lowest validation perplexity.",
    )
    table(
        doc,
        ["Hyperparameter", "Selected value", "Meaning"],
        [
            ["Lambdas", METRICS["best_lambdas"], "40% unigram, 30% bigram, 20% trigram, 10% 4-gram"],
            ["k", METRICS["best_k"], "Very small smoothing, suitable for this sparse n-gram setting"],
            ["Validation perplexity", METRICS["best_valid_pp"], "Best LM2 validation score found by grid search"],
        ],
    )
    paragraph(
        doc,
        "The selected lambda values show that lower-order n-grams are more useful for this corpus. This is reasonable because exact 4-gram patterns are still sparse, even after increasing the corpus size.",
    )

    heading(doc, "5. Main Evaluation Result", 1)
    table(
        doc,
        ["Model", "Method", "Test perplexity"],
        [
            ["LM1", "Unsmoothed backoff 4-gram", METRICS["lm1_test_pp"]],
            ["LM2", "Interpolation with add-k smoothing", METRICS["lm2_test_pp"]],
        ],
    )
    doc.add_picture(str(model_img), width=Inches(5.8))
    paragraph(
        doc,
        "Lower perplexity means the model predicts the test data better. LM1 achieves a test perplexity of 27.2196, while LM2 achieves 45.7879. Therefore, LM1 performs better in this experiment. This does not mean interpolation is always worse; it means that for this Khmer corpus, vocabulary setting, and tokenizer, backoff handled the sparse 4-gram patterns more effectively.",
    )

    heading(doc, "6. Vocabulary Size Experiment", 1)
    paragraph(
        doc,
        "I also tested different vocabulary sizes because vocabulary size has a strong effect on both perplexity and readability. A smaller vocabulary converts more rare words to <UNK>, which can reduce perplexity but makes generated text less readable. A larger vocabulary keeps more real words, but it increases sparsity and makes prediction harder.",
    )
    table(
        doc,
        ["Vocab cap", "Actual vocab", "Test <UNK>", "LM1 PP", "LM2 PP"],
        [(row[0], row[1], row[3], row[4], row[5]) for row in VOCAB_ROWS],
    )
    doc.add_picture(str(vocab_img), width=Inches(5.8))
    paragraph(
        doc,
        "The graph shows that perplexity increases as vocabulary size increases. This happens because the model has more possible next tokens to choose from and many tokens appear only a few times. I selected vocabulary size 500 as a balanced setting: it keeps more real Khmer words than 100 while avoiding the extreme sparsity of 1000.",
    )

    heading(doc, "7. Text Generation", 1)
    paragraph(
        doc,
        "The text generator starts from a seed phrase and predicts the next token repeatedly. For presentation, LM1 is recommended because it has lower test perplexity and usually produces cleaner short output. A length of 20 to 30 tokens is easier to read than a very long generation.",
    )
    table(
        doc,
        ["Seed", "Model", "Generated example"],
        [
            ["ប្រាសាទ", "LM1 Backoff", "ប្រាសាទ នេះ កសាង នៅ ចុង សតវត្សរ៍ ទី ១២ និង ១៣។"],
            ["ប្រាសាទ", "LM2 Interpolation", "ប្រាសាទ នេះ បង្ហាញ ឲ្យដឹង ថា ត្រូវបាន ទៅជា <UNK> ទើប នៅក្នុង រូបរាង <UNK> <UNK> របស់ ព្រះអង្គ <UNK>។"],
        ],
    )
    paragraph(
        doc,
        "The generated text is not perfect, which is expected for an n-gram model. N-gram models do not understand meaning like modern neural language models; they generate text by reusing statistical patterns from the corpus.",
    )

    heading(doc, "8. Conclusion", 1)
    paragraph(
        doc,
        "This project satisfies the required tasks: corpus selection, 70/10/20 data splitting, tokenization, vocabulary limiting with <UNK>, two 4-gram language models, validation-based tuning, perplexity evaluation, and text generation. The strongest result is that LM1 Backoff performs better than LM2 Interpolation on the final test set.",
    )
    paragraph(
        doc,
        "The most important lesson from the experiment is the trade-off between perplexity and readability. A small vocabulary gives lower perplexity but more <UNK> tokens, while a larger vocabulary produces more readable text but increases sparsity and perplexity. The project also shows that Khmer tokenization quality matters, so using a Khmer-specific tokenizer improves the linguistic quality of the input tokens.",
    )

    heading(doc, "References", 1)
    for url in SOURCE_URLS:
        paragraph(doc, url, size=9)

    doc.save(REPORT_PATH)


if __name__ == "__main__":
    build()
    print(REPORT_PATH)
