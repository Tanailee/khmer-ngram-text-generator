from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT_DIR = Path(__file__).resolve().parent
OUT_DIR = PROJECT_DIR / "report_assets"
OUT_DIR.mkdir(exist_ok=True)
REPORT_PATH = PROJECT_DIR / "NLP_Mini_Project_1_Final_Report.docx"


METRICS = {
    "articles": 5,
    "corpus_chars": 96199,
    "tokenizer": "khmernltk.word_tokenize",
    "total_tokens": 19681,
    "train_tokens": 13776,
    "valid_tokens": 1968,
    "test_tokens": 3937,
    "vocab_size": 500,
    "valid_unk": 660,
    "test_unk": 1433,
    "best_lambdas": "(0.4, 0.3, 0.2, 0.1)",
    "best_k": 0.001,
    "best_valid_pp": 47.4493,
    "lm1_test_pp": 27.2196,
    "lm2_test_pp": 45.7879,
}

VOCAB_ROWS = [
    (100, 100, 1164, 2344, 7.0151, 9.2468, "(0.4, 0.3, 0.2, 0.1)", 0.001),
    (200, 200, 930, 2021, 11.3045, 16.1981, "(0.4, 0.3, 0.2, 0.1)", 0.001),
    (300, 300, 808, 1783, 16.2645, 24.7400, "(0.4, 0.3, 0.2, 0.1)", 0.001),
    (500, 500, 660, 1433, 27.2196, 45.7879, "(0.4, 0.3, 0.2, 0.1)", 0.001),
    (700, 700, 574, 1233, 38.2488, 68.3284, "(0.4, 0.3, 0.2, 0.1)", 0.001),
    (1000, 1000, 523, 1119, 49.2772, 92.2916, "(0.4, 0.3, 0.2, 0.1)", 0.001),
]

SOURCE_URLS = [
    "https://km.wikipedia.org/wiki/ប្រាសាទ_អង្គរវត្ត",
    "https://km.wikipedia.org/wiki/អង្គរធំ",
    "https://km.wikipedia.org/wiki/ប្រាសាទបាយ័ន",
    "https://km.wikipedia.org/wiki/ខេត្តសៀមរាប",
    "https://km.wikipedia.org/wiki/ចក្រភពខ្មែរ",
]

LM1_SAMPLE = (
    "ប្រាសាទ នេះ កសាង នៅ ចុង សតវត្សរ៍ ទី ១២ និង ១៣។"
)
LM2_SAMPLE = (
    "ប្រាសាទ នេះ បង្ហាញ ឲ្យដឹង ថា ត្រូវបាន ទៅជា <UNK> ទើប នៅក្នុង រូបរាង "
    "<UNK> <UNK> របស់ ព្រះអង្គ <UNK>។"
)


def save_charts():
    split_path = OUT_DIR / "dataset_split.png"
    model_path = OUT_DIR / "model_perplexity.png"
    vocab_path = OUT_DIR / "vocab_perplexity.png"

    plt.figure(figsize=(7, 4))
    labels = ["Training", "Validation", "Testing"]
    counts = [METRICS["train_tokens"], METRICS["valid_tokens"], METRICS["test_tokens"]]
    bars = plt.bar(labels, counts, color=["#2E86AB", "#F6C85F", "#6F4E7C"])
    plt.title("Corpus Split: Training, Validation, Testing")
    plt.ylabel("Number of Tokens")
    plt.grid(axis="y", alpha=0.25)
    for bar, count in zip(bars, counts):
        plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 100, str(count), ha="center")
    plt.tight_layout()
    plt.savefig(split_path, dpi=180)
    plt.close()

    plt.figure(figsize=(7, 4))
    bars = plt.bar(
        ["LM1 Backoff", "LM2 Interpolation"],
        [METRICS["lm1_test_pp"], METRICS["lm2_test_pp"]],
        color=["#2E86AB", "#D1495B"],
    )
    plt.title("Test Perplexity Comparison")
    plt.ylabel("Perplexity (Lower is Better)")
    plt.grid(axis="y", alpha=0.25)
    for bar, value in zip(bars, [METRICS["lm1_test_pp"], METRICS["lm2_test_pp"]]):
        plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.8, f"{value:.4f}", ha="center")
    plt.tight_layout()
    plt.savefig(model_path, dpi=180)
    plt.close()

    caps = [row[0] for row in VOCAB_ROWS]
    lm1 = [row[4] for row in VOCAB_ROWS]
    lm2 = [row[5] for row in VOCAB_ROWS]
    plt.figure(figsize=(8, 4.5))
    plt.plot(caps, lm1, marker="o", linewidth=2, label="LM1 Backoff")
    plt.plot(caps, lm2, marker="o", linewidth=2, label="LM2 Interpolation")
    plt.title("Vocabulary Size vs Test Perplexity")
    plt.xlabel("Vocabulary Cap")
    plt.ylabel("Test Perplexity (Lower is Better)")
    plt.grid(alpha=0.25)
    plt.legend()
    plt.tight_layout()
    plt.savefig(vocab_path, dpi=180)
    plt.close()

    return split_path, model_path, vocab_path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(run, size=10, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans Khmer")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc, text, size=10.5, bold=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    set_font(r, size=14 if level == 1 else 12, bold=True)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = str(header)
        set_cell_shading(hdr[i], "D9EAF7")
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=9)
    return table


def build_report():
    split_path, model_path, vocab_path = save_charts()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Mini Project 1: Text Generation with 4-gram Language Models")
    set_font(r, size=16, bold=True)
    add_paragraph(doc, "Natural Language Processing | Master of Data Science", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Corpus: Combined Khmer Wikipedia articles related to Angkor Wat", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "1. Objective", 1)
    add_paragraph(
        doc,
        "This project builds and evaluates two 4-gram language models for Khmer text generation. "
        "LM1 uses an unsmoothed backoff method, while LM2 uses interpolation with add-k smoothing. "
        "The models are evaluated with perplexity and demonstrated using generated text examples.",
    )

    add_heading(doc, "2. Corpus and Preprocessing", 1)
    add_paragraph(
        doc,
        "To improve corpus size, five related Khmer Wikipedia articles were combined. The notebook removes noisy English-only lines, "
        "removes artificial article markers, separates Khmer punctuation, and uses khmernltk.word_tokenize for Khmer word segmentation. "
        "Vocabulary is limited and remaining tokens are replaced with <UNK>, matching the assignment requirement.",
    )
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Articles in corpus", METRICS["articles"]],
            ["Corpus characters", f"{METRICS['corpus_chars']:,}"],
            ["Tokenizer", METRICS["tokenizer"]],
            ["Total tokens", f"{METRICS['total_tokens']:,}"],
            ["Training / Validation / Testing", f"{METRICS['train_tokens']:,} / {METRICS['valid_tokens']:,} / {METRICS['test_tokens']:,}"],
            ["Vocabulary size", METRICS["vocab_size"]],
            ["Validation <UNK> / Test <UNK>", f"{METRICS['valid_unk']} / {METRICS['test_unk']}"],
        ],
    )
    doc.add_paragraph()
    doc.add_picture(str(split_path), width=Inches(5.8))

    add_heading(doc, "3. Model Design", 1)
    add_paragraph(
        doc,
        "LM1 Backoff first tries the 4-gram probability. If the 4-gram is unseen, it backs off to trigram, bigram, and unigram probabilities. "
        "LM2 Interpolation combines unigram, bigram, trigram, and 4-gram probabilities using lambda weights. Each LM2 term uses add-k smoothing."
    )
    add_table(
        doc,
        ["Model", "Method", "Key idea"],
        [
            ["LM1", "Unsmoothed backoff 4-gram", "Use longest available context; back off when unseen."],
            ["LM2", "Interpolation + add-k smoothing", "Combine all n-gram levels and smooth unseen events."],
        ],
    )

    add_heading(doc, "4. Hyperparameter Tuning", 1)
    add_paragraph(
        doc,
        "Grid search was used for LM2. Several lambda combinations and k values were tested on the validation set. "
        "The best setting was selected by lowest validation perplexity, then final performance was measured on the test set."
    )
    add_table(
        doc,
        ["Best LM2 parameter", "Value"],
        [
            ["Best lambdas", METRICS["best_lambdas"]],
            ["Best k", METRICS["best_k"]],
            ["Validation perplexity", METRICS["best_valid_pp"]],
        ],
    )

    add_heading(doc, "5. Evaluation Results", 1)
    add_table(
        doc,
        ["Model", "Method", "Best parameters", "Test perplexity"],
        [
            ["LM1", "Unsmoothed backoff 4-gram", "No lambda/k", METRICS["lm1_test_pp"]],
            ["LM2", "Interpolation + add-k smoothing", f"lambdas={METRICS['best_lambdas']}, k={METRICS['best_k']}", METRICS["lm2_test_pp"]],
        ],
    )
    add_paragraph(
        doc,
        "Lower perplexity is better. LM1 Backoff performs better than LM2 Interpolation in this experiment. "
        "The best LM2 lambda values give more weight to lower-order n-grams, showing that exact 4-gram patterns are still sparse.",
    )
    doc.add_picture(str(model_path), width=Inches(5.8))

    add_heading(doc, "6. Vocabulary Size Experiment", 1)
    add_paragraph(
        doc,
        "Vocabulary size controls the trade-off between readability and perplexity. A small vocabulary creates many <UNK> tokens but can lower perplexity. "
        "A larger vocabulary keeps more real Khmer words but makes n-gram counts more sparse."
    )
    add_table(
        doc,
        ["Vocab", "Test <UNK>", "LM1 PP", "LM2 PP"],
        [(row[0], row[3], row[4], row[5]) for row in VOCAB_ROWS],
    )
    doc.add_picture(str(vocab_path), width=Inches(5.8))

    add_heading(doc, "7. Text Generation Examples", 1)
    add_paragraph(doc, "Seed phrase: ប្រាសាទ", bold=True)
    add_paragraph(doc, "LM1 Backoff generated text:", bold=True)
    add_paragraph(doc, LM1_SAMPLE)
    add_paragraph(doc, "LM2 Interpolation generated text:", bold=True)
    add_paragraph(doc, LM2_SAMPLE)
    add_paragraph(
        doc,
        "For presentation, LM1 Backoff is recommended because it has lower perplexity and usually produces cleaner short generated text. "
        "A generation length of 20-30 tokens is easier to read.",
    )

    add_heading(doc, "8. Conclusion", 1)
    add_paragraph(
        doc,
        "The project satisfies the assignment requirements: corpus selection, 70/10/20 split, tokenization, vocabulary limiting with <UNK>, "
        "two 4-gram language models, LM2 hyperparameter tuning, perplexity evaluation, and text generation. "
        "LM1 Backoff achieved the best test perplexity. LM2 is smoother but performed worse on this corpus. "
        "The main limitations are n-gram sparsity, the vocabulary/readability trade-off, and imperfect naturalness of generated Khmer text.",
    )

    add_heading(doc, "References", 1)
    for url in SOURCE_URLS:
        add_paragraph(doc, url, size=9)

    doc.save(REPORT_PATH)


if __name__ == "__main__":
    build_report()
    print(REPORT_PATH)
